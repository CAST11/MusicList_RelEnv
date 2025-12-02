"""
Script to generate a txt file that contains the names of all the songs contained in 'MiBotiquin2024' folder
Contributor:
    Celina Soto
"""
# Import required modules
import os
import sys 
from pathlib import Path
from tinytag import TinyTag
import datetime
from tkinter import *
from tkinter import filedialog
from tkinter import ttk
import pandas as pd
import openpyxl
from docx import Document


class AudioFilesList:
    ''' Description goes here'''

    def __init__(self):
        self.audio_directory: str = None
        self.library_path_name: str = None
        self.library_filename: str = None
        self.library_path: str = None
        self.selected_ext: str = None
        self.label_file_explorer: Label = None
        self.button_generate_list: Button = None
        self.songs_title_list_generated: bool = False
        self.button_exit: Button = None
        self.song_path: Path = None
        self.song_title: str = None
        self.songs_title_list: list = []       

    # Function for opening the file explorer window and selecting audio folder
    def browseMusicFolder(self):        
        folder = filedialog.askdirectory(initialdir = "/", title = "Select a Folder")
        if folder:
            self.audio_directory = folder
            if hasattr(self, "button_generate_list") and self.button_generate_list:
                self.button_generate_list.config(state=NORMAL)
            if hasattr(self, "label_music_directory_explorer") and self.label_music_directory_explorer:
                # Change label contents
                self.label_music_directory_explorer.configure(text="Directory Selected: "+ self.audio_directory)
        return folder


    # Function for opening the file explorer window and selecting folder for library file
    def saveLibraryFile(self):  
        file_selected = filedialog.asksaveasfilename(filetypes=[('Text Files', '*.txt'), ('Excel Files', '*.xlsx'), ('Word Files', '*.docx'), ('All Files', '*.*')], defaultextension='.txt') 
        
        # User pressed cancel
        if not file_selected:
            print("No file selected - Operation cancelled")
            return 
        
        # User selected a file
        self.library_path_name = file_selected
        print(f"self.library_path_name: {self.library_path_name}")

        # Get the selected filetype based on the extension in the dropdown
        self.selected_ext = os.path.splitext(self.library_path_name)
        print(f"self.selected_ext[extension]: {self.selected_ext[1]}")

        # Get the filename and path for the library 
        print(f"self.selected_ext[root]: {self.selected_ext[0]}")

        # Get the selected library directory 
        self.library_path = os.path.dirname(self.library_path_name)
        print(f"self.library_path: {self.library_path}")
        #print(f"self.selected_ext[root]: {self.selected_ext[0]}")

        # Get the selected filename assigned by user 
        self.library_filename = os.path.basename(self.library_path_name)
        print(f"self.library_filename: {self.library_filename}")

        # Change label contents
        # Update UI label safely
        if self.label_file_explorer:
            self.label_file_explorer.configure(text=f"File Opened: {self.library_path_name}")
        
        

    def create_txt_file(self):
        if not self.library_path_name:
            raise ValueError("library_path_name is not set before calling create_txt_file()")
        
        with open (self.library_path_name, mode="w", encoding='utf-8') as file:
            file.write(f"File songs library \n")
            for i, song in enumerate (self.songs_title_list, start=1):
                print(f"Song title: {song}") 
                file.write(f"{i}: {song} \n")
                         

    def create_excel_file(self):
        headers = ["File songs library"]

        # Convert list of lists to DataFrame with headers
        df = pd.DataFrame(self.songs_title_list, columns=headers)
        df.index = range(1, len(df) + 1)
        with pd.ExcelWriter(self.library_path_name, engine="openpyxl", mode="w") as writer:
            df.to_excel(writer, sheet_name="SongsList", index=True)


    def create_word_file(self):
        doc = Document()
                     
        # Create new document    
        doc.add_heading("File songs library", level=1)
        for i, song in enumerate (self.songs_title_list, start=1):
            doc.add_paragraph(f"{i}: {song} \n")

        doc.save(self.library_path_name)  


    # Function to create song list 
    def create_songs_list(self):    
        if not self.songs_title_list_generated:    
        # Assign directory
        #c:\Users\dona_\Documents\Python\Musicoterapia\MiBotiquinLista
            print(f"main directory: {self.audio_directory}")

            log_file_path = os.path.join(self.audio_directory, f"Error_Exec_Log_{datetime.date.today()}.txt") 
            print(f"logs file path: {log_file_path}")
    
            for root, dirs, files in os.walk(self.audio_directory, topdown=True):

                for file in files:
                    self.song_path = os.path.join(root, file)
                    #print(f"song path: {self.song_path}")
            
                    try:                    
                        #Read audio file metadata: Title
                        tag: TinyTag = TinyTag.get(self.song_path)
                        self.song_title = tag.title                    
                        self.songs_title_list.append(self.song_title)                    
                        
                    except Exception as Exc:
                        print(f"Exc: {Exc}")
                        # Create Error_Exec_Log_ActualDate.txt filename
                        with open (f"{log_file_path}", mode='a') as file:
                            file.write(f"File: {self.song_path}\n")
                            print(f"{self.song_path} is not am mp3 file")
            self.songs_title_list_generated = True
 

    def create_song_file(self):
        self.create_songs_list()

        if self.selected_ext[1] == ".txt":
            self.create_txt_file()                                           
                
        elif self.selected_ext[1] == ".xlsx":
            self.create_excel_file()
                    
        elif self.selected_ext[1] == ".docx":
            self.create_word_file()

        else:
            print("Fileformat is not supported")            
            

    def run(self):
        # Create the uw window
        uw = Tk()

        # Set user window title 
        uw.title("Welcome to your rescue")

        # Set window size (width x height)
        uw.geometry('560x300')

        # Set window background color
        uw.config(background = "gray")

        # Create a File Explorer label
        self.label_music_directory_explorer = Label(uw,
                            text="Directory Explorer - Audio",
                            width=75, height=2,
                            fg="blue")
        
        # Create a File Explorer label
        self.label_file_explorer = Label(uw, 
                            text = "File Saving *- Library",
                            width = 75, height = 2, 
                            fg = "blue")
        
        # Create button to select audio folder directory
        self.button_dir_audio_explore = Button(uw,
                        text="Browse audio files directory",
                        command=self.browseMusicFolder)
        
        # Create button to select filename for library and directory
        self.button_dir_filename_for_library = Button(uw,
                        text="Provide file name",
                        command=self.saveLibraryFile)

        # Create button to generate the library file
        self.button_generate_list = Button(uw,
                        text="Generate list of files",
                        command=self.create_song_file,
                        state=DISABLED)
        
        # Exit button 
        self.button_exit = Button(uw,
                     text="Exit",
                     command=exit)

        # Grid method is chosen for placing the widgets at respective positions 
        # in a table like structure by specifying rows and columns
        self.label_music_directory_explorer.grid(column = 1, row = 1, padx= 15, pady=10)

        self.label_file_explorer.grid(column = 1, row = 2, padx= 15, pady=10)

        self.button_dir_audio_explore.grid(column=1, row=4, padx= 10, pady=10)

        self.button_dir_filename_for_library.grid(column=1, row=5, padx= 10, pady=10)


        self.button_generate_list.grid(column=1, row=7, pady=10)

        self.button_exit.grid(column=1, row=10, pady=10)

        uw.mainloop()

if __name__ =="__main__":
    FilesList = AudioFilesList()
    FilesList.run()